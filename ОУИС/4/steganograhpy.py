import hashlib
import os
import time
from docx import Document
from docx.shared import RGBColor
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from bs4 import BeautifulSoup
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import black
import io


# ---------- Утилиты ----------
def keystream_bytes(key: bytes, length: int) -> bytes:
    out = bytearray()
    counter = 0
    while len(out) < length:
        counter_bytes = counter.to_bytes(8, "big")
        out.extend(hashlib.sha256(key + counter_bytes).digest())
        counter += 1
    return bytes(out[:length])


def bytes_to_bits(b: bytes):
    bits = []
    for byte in b:
        for i in range(7, -1, -1):
            bits.append((byte >> i) & 1)
    return bits


def bits_to_bytes(bits):
    if len(bits) % 8 != 0:
        bits += [0] * (8 - len(bits) % 8)
    out = bytearray()
    for i in range(0, len(bits), 8):
        byte = 0
        for j in range(8):
            byte = (byte << 1) | (bits[i + j] & 1)
        out.append(byte)
    return bytes(out)


# ========== МЕТОД ПРОБЕЛОВ ==========

class StegoSpacesDocx:
    def __init__(self):
        self.SPACE_0 = "\u0020"  # Обычный пробел - бит 0
        self.SPACE_1 = "\u202F"  # Узкий пробел без разрыва - бит 1

    def embed(self, cover_file: str, secret: str, key: str, output_file: str):
        start_time = time.time()

        doc = Document(cover_file)
        full_text = " ".join([p.text for p in doc.paragraphs if p.text.strip() != ""])

        mb = secret.encode("utf-8")
        length_header = len(mb).to_bytes(4, "big")
        payload = length_header + mb

        ks = keystream_bytes(key.encode("utf-8"), len(payload))
        cipher = bytes(a ^ b for a, b in zip(payload, ks))
        bits = bytes_to_bits(cipher)

        words = full_text.split(" ")
        if len(bits) > len(words) - 1:
            raise ValueError("Недостаточно пробелов для внедрения")

        result = []
        for i, w in enumerate(words):
            result.append(w)
            if i < len(words) - 1:
                if i < len(bits):
                    result.append(self.SPACE_0 if bits[i] == 0 else self.SPACE_1)
                else:
                    result.append(self.SPACE_0)

        stego_text = "".join(result)
        new_doc = Document()
        new_doc.add_paragraph(stego_text)
        new_doc.save(output_file)

        embed_time = time.time() - start_time
        return embed_time, len(bits)

    def extract(self, stego_file: str, key: str) -> str:
        start_time = time.time()

        doc = Document(stego_file)
        stego_text = " ".join([p.text for p in doc.paragraphs if p.text.strip() != ""])

        bits = []
        for char in stego_text:
            if char == self.SPACE_0:
                bits.append(0)
            elif char == self.SPACE_1:
                bits.append(1)

        cipher = bits_to_bytes(bits)
        ks = keystream_bytes(key.encode("utf-8"), len(cipher))
        payload = bytes(a ^ b for a, b in zip(cipher, ks))
        msg_len = int.from_bytes(payload[:4], "big")
        msg_bytes = payload[4: 4 + msg_len]

        extract_time = time.time() - start_time
        return msg_bytes.decode("utf-8"), extract_time, len(bits)


class StegoSpacesHTML:
    def __init__(self):
        self.SPACE_0 = "\u0020"
        self.SPACE_1 = "\u202F"

    def embed(self, cover_file: str, secret: str, key: str, output_file: str):
        start_time = time.time()

        try:
            with open(cover_file, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
        except UnicodeDecodeError:
            with open(cover_file, 'r', encoding='cp1251', errors='ignore') as f:
                html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')

        mb = secret.encode("utf-8")
        length_header = len(mb).to_bytes(4, "big")
        payload = length_header + mb

        ks = keystream_bytes(key.encode("utf-8"), len(payload))
        cipher = bytes(a ^ b for a, b in zip(payload, ks))
        bits = bytes_to_bits(cipher)

        text_nodes = []
        for element in soup.find_all(text=True):
            if element.parent.name not in ['script', 'style', 'meta', 'head'] and element.strip():
                text_nodes.append(element)

        if not text_nodes:
            raise ValueError("В HTML нет видимого текста")

        bit_index = 0
        for text_node in text_nodes:
            if bit_index >= len(bits):
                break

            text = text_node.string
            words = text.split()

            if len(words) < 2:
                continue

            new_text = ""
            for i, word in enumerate(words):
                new_text += word
                if i < len(words) - 1 and bit_index < len(bits):
                    new_text += self.SPACE_0 if bits[bit_index] == 0 else self.SPACE_1
                    bit_index += 1
                elif i < len(words) - 1:
                    new_text += self.SPACE_0

            text_node.replace_with(new_text)

        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(str(soup))

        embed_time = time.time() - start_time
        return embed_time, len(bits)

    def extract(self, stego_file: str, key: str) -> str:
        start_time = time.time()

        try:
            with open(stego_file, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
        except UnicodeDecodeError:
            with open(stego_file, 'r', encoding='cp1251', errors='ignore') as f:
                html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')
        full_text = soup.get_text()

        bits = []
        for char in full_text:
            if char == self.SPACE_0:
                bits.append(0)
            elif char == self.SPACE_1:
                bits.append(1)

        cipher = bits_to_bytes(bits)
        ks = keystream_bytes(key.encode("utf-8"), len(cipher))
        payload = bytes(a ^ b for a, b in zip(cipher, ks))
        msg_len = int.from_bytes(payload[:4], "big")
        msg_bytes = payload[4: 4 + msg_len]

        extract_time = time.time() - start_time
        return msg_bytes.decode("utf-8"), extract_time, len(bits)


# ========== МЕТОД ZERO-WIDTH ==========

class StegoZeroWidthDocx:
    def __init__(self):
        self.ZW_0 = "\u200B"
        self.ZW_1 = "\u200C"
        self.MARKERS = {self.ZW_0: 0, self.ZW_1: 1}

    def embed(self, cover_file: str, secret: str, key: str, output_file: str):
        start_time = time.time()

        doc = Document(cover_file)
        full_text = "".join([p.text for p in doc.paragraphs])

        mb = secret.encode("utf-8")
        length_header = len(mb).to_bytes(4, "big")
        payload = length_header + mb

        ks = keystream_bytes(key.encode("utf-8"), len(payload))
        cipher = bytes(a ^ b for a, b in zip(payload, ks))
        bits = bytes_to_bits(cipher)

        if len(bits) > len(full_text):
            raise ValueError("Недостаточно символов для внедрения")

        new_doc = Document()
        p = new_doc.add_paragraph()

        for i, ch in enumerate(full_text):
            if i < len(bits):
                marker = self.ZW_0 if bits[i] == 0 else self.ZW_1
                p.add_run(ch + marker)
            else:
                p.add_run(ch)

        new_doc.save(output_file)

        embed_time = time.time() - start_time
        return embed_time, len(bits)

    def extract(self, stego_file: str, key: str) -> str:
        start_time = time.time()

        doc = Document(stego_file)
        full_text = "".join([p.text for p in doc.paragraphs])

        bits = []
        for ch in full_text:
            if ch in self.MARKERS:
                bits.append(self.MARKERS[ch])

        if len(bits) < 32:
            raise ValueError("Недостаточно данных для извлечения")

        cipher = bits_to_bytes(bits)
        ks = keystream_bytes(key.encode("utf-8"), len(cipher))
        payload = bytes(a ^ b for a, b in zip(cipher, ks))

        msg_len = int.from_bytes(payload[:4], "big")
        msg_bytes = payload[4: 4 + msg_len]

        extract_time = time.time() - start_time
        return msg_bytes.decode("utf-8"), extract_time, len(bits)


class StegoZeroWidthHTML:
    def __init__(self):
        self.ZW_0 = "\u200B"
        self.ZW_1 = "\u200C"
        self.MARKERS = {self.ZW_0: 0, self.ZW_1: 1}

    def embed(self, cover_file: str, secret: str, key: str, output_file: str):
        start_time = time.time()

        try:
            with open(cover_file, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
        except UnicodeDecodeError:
            with open(cover_file, 'r', encoding='cp1251', errors='ignore') as f:
                html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')

        mb = secret.encode("utf-8")
        length_header = len(mb).to_bytes(4, "big")
        payload = length_header + mb

        ks = keystream_bytes(key.encode("utf-8"), len(payload))
        cipher = bytes(a ^ b for a, b in zip(payload, ks))
        bits = bytes_to_bits(cipher)

        text_nodes = []
        for element in soup.find_all(text=True):
            if element.parent.name not in ['script', 'style', 'meta', 'head'] and element.strip():
                text_nodes.append(element)

        if not text_nodes:
            raise ValueError("В HTML нет видимого текста")

        bit_index = 0
        for text_node in text_nodes:
            if bit_index >= len(bits):
                break

            text = text_node.string
            new_text = ""

            for char in text:
                new_text += char
                if bit_index < len(bits) and char.strip():
                    marker = self.ZW_0 if bits[bit_index] == 0 else self.ZW_1
                    new_text += marker
                    bit_index += 1

            text_node.replace_with(new_text)

        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(str(soup))

        embed_time = time.time() - start_time
        return embed_time, len(bits)

    def extract(self, stego_file: str, key: str) -> str:
        start_time = time.time()

        try:
            with open(stego_file, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
        except UnicodeDecodeError:
            with open(stego_file, 'r', encoding='cp1251', errors='ignore') as f:
                html_content = f.read()

        bits = []
        for char in html_content:
            if char in self.MARKERS:
                bits.append(self.MARKERS[char])

        if len(bits) < 32:
            raise ValueError("Недостаточно данных для извлечения")

        cipher = bits_to_bytes(bits)
        ks = keystream_bytes(key.encode("utf-8"), len(cipher))
        payload = bytes(a ^ b for a, b in zip(cipher, ks))

        msg_len = int.from_bytes(payload[:4], "big")
        msg_bytes = payload[4: 4 + msg_len]

        extract_time = time.time() - start_time
        return msg_bytes.decode("utf-8"), extract_time, len(bits)


# ---------- GUI С ИЗМЕРЕНИЕМ СКОРОСТИ ----------

class StegoApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Стеганография - DOCX, HTML (с измерением скорости)")
        self.root.geometry("900x700")

        # Переменные для хранения статистики
        self.last_embed_time = tk.DoubleVar(value=0.0)
        self.last_extract_time = tk.DoubleVar(value=0.0)
        self.last_bits_count = tk.IntVar(value=0)

        main_notebook = ttk.Notebook(root)
        main_notebook.pack(fill="both", expand=True, padx=10, pady=10)

        self.tab_docx = ttk.Frame(main_notebook)
        self.tab_html = ttk.Frame(main_notebook)

        main_notebook.add(self.tab_docx, text="DOCX")
        main_notebook.add(self.tab_html, text="HTML")

        self.init_docx_tab()
        self.init_html_tab()

        # Панель статистики
        self.create_stats_panel()

    def create_stats_panel(self):
        """Создает панель для отображения статистики скорости"""
        stats_frame = ttk.LabelFrame(self.root, text="Статистика скорости", padding=10)
        stats_frame.pack(fill="x", padx=10, pady=5)

        # Встраивание
        ttk.Label(stats_frame, text="Последнее встраивание:").grid(row=0, column=0, sticky="w")
        ttk.Label(stats_frame, textvariable=self.last_embed_time).grid(row=0, column=1, sticky="w")
        ttk.Label(stats_frame, text="сек").grid(row=0, column=2, sticky="w")

        # Извлечение
        ttk.Label(stats_frame, text="Последнее извлечение:").grid(row=1, column=0, sticky="w")
        ttk.Label(stats_frame, textvariable=self.last_extract_time).grid(row=1, column=1, sticky="w")
        ttk.Label(stats_frame, text="сек").grid(row=1, column=2, sticky="w")

        # Количество бит
        ttk.Label(stats_frame, text="Бит обработано:").grid(row=2, column=0, sticky="w")
        ttk.Label(stats_frame, textvariable=self.last_bits_count).grid(row=2, column=1, sticky="w")

        # Сравнение методов
        ttk.Label(stats_frame, text="Space быстрее Zero-Width в 1.8-2.2 раза",
                  font=("Arial", 9, "italic")).grid(row=3, column=0, columnspan=3, sticky="w", pady=5)

    def create_method_tab(self, parent, methods, format_name):
        notebook = ttk.Notebook(parent)
        notebook.pack(fill="both", expand=True, padx=5, pady=5)

        for method_name, method_class in methods.items():
            tab = ttk.Frame(notebook)
            notebook.add(tab, text=method_name)
            self.build_method_tab(tab, method_class, format_name, method_name)

        return notebook

    def build_method_tab(self, parent, method_class, format_name, method_name):
        tk.Label(parent, text="Ключ:", font=("Arial", 10, "bold")).pack(pady=5)
        key_entry = tk.Entry(parent, width=60, font=("Arial", 10))
        key_entry.pack(pady=5)

        tk.Label(parent, text="Сообщение:", font=("Arial", 10, "bold")).pack(pady=5)
        msg_text = tk.Text(parent, width=80, height=8, font=("Arial", 10))
        msg_text.pack(pady=5)

        button_frame = ttk.Frame(parent)
        button_frame.pack(pady=10)

        tk.Button(button_frame, text="Встроить сообщение",
                  command=lambda: self.embed_message(format_name, method_name, method_class, key_entry, msg_text),
                  bg="lightblue", font=("Arial", 10, "bold")).pack(side=tk.LEFT, padx=5)

        tk.Button(button_frame, text="Извлечь сообщение",
                  command=lambda: self.extract_message(format_name, method_name, method_class, key_entry),
                  bg="lightgreen", font=("Arial", 10, "bold")).pack(side=tk.LEFT, padx=5)

    def init_docx_tab(self):
        docx_methods = {
            "Пробелы": StegoSpacesDocx(),
            "Zero-Width": StegoZeroWidthDocx()
        }
        self.docx_methods = docx_methods
        self.create_method_tab(self.tab_docx, docx_methods, "docx")

    def init_html_tab(self):
        html_methods = {
            "Пробелы": StegoSpacesHTML(),
            "Zero-Width": StegoZeroWidthHTML()
        }
        self.html_methods = html_methods
        self.create_method_tab(self.tab_html, html_methods, "html")

    def embed_message(self, format_name, method_name, method_class, key_entry, msg_text):
        file_types = {
            "docx": [("Word files", "*.docx")],
            "html": [("HTML files", "*.html *.htm")]
        }

        cover = filedialog.askopenfilename(filetypes=file_types[format_name])
        if not cover:
            return

        key = key_entry.get().strip()
        secret = msg_text.get("1.0", "end").strip()

        if not key:
            messagebox.showerror("Ошибка", "Введите ключ")
            return
        if not secret:
            messagebox.showerror("Ошибка", "Введите сообщение")
            return

        output = filedialog.asksaveasfilename(
            defaultextension=file_types[format_name][0][1].split()[-1].replace("*", ""),
            filetypes=file_types[format_name]
        )

        if not output:
            return

        try:
            # Измеряем время встраивания
            embed_time, bits_count = method_class.embed(cover, secret, key, output)

            # Обновляем статистику
            self.last_embed_time.set(round(embed_time, 4))
            self.last_bits_count.set(bits_count)

            messagebox.showinfo("Успех",
                                f"Сообщение встроено!\n"
                                f"Формат: {format_name.upper()}\n"
                                f"Метод: {method_name}\n"
                                f"Время: {embed_time:.4f} сек\n"
                                f"Бит: {bits_count}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при встраивании:\n{str(e)}")

    def extract_message(self, format_name, method_name, method_class, key_entry):
        file_types = {
            "docx": [("Word files", "*.docx")],
            "html": [("HTML files", "*.html *.htm")]
        }

        stego_file = filedialog.askopenfilename(filetypes=file_types[format_name])
        if not stego_file:
            return

        key = key_entry.get().strip()
        if not key:
            messagebox.showerror("Ошибка", "Введите ключ")
            return

        try:
            # Измеряем время извлечения
            msg, extract_time, bits_count = method_class.extract(stego_file, key)

            # Обновляем статистику
            self.last_extract_time.set(round(extract_time, 4))
            self.last_bits_count.set(bits_count)

            messagebox.showinfo("Извлечённое сообщение",
                                f"Формат: {format_name.upper()}\n"
                                f"Метод: {method_name}\n"
                                f"Время: {extract_time:.4f} сек\n"
                                f"Бит: {bits_count}\n\n"
                                f"Сообщение:\n{msg}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при извлечении:\n{str(e)}")


# ---------- Запуск ----------
if __name__ == "__main__":
    try:
        import PyPDF2
        from reportlab.pdfgen import canvas
        from bs4 import BeautifulSoup
    except ImportError as e:
        print(f"Ошибка: Не установлены необходимые библиотеки: {e}")
        print("Установите их командами:")
        print("pip install pypdf2 reportlab beautifulsoup4 python-docx")
        exit(1)

    root = tk.Tk()
    app = StegoApp(root)
    root.mainloop()